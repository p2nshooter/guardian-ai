# ==============================================================================
# Copyright (c) 2024-2026 Yusron Efendi. All rights reserved.
# Platform Architecture: AXTO (axto.io) — Sovereign AI Infrastructure
# Author & Architect: Yusron Efendi <hallo@axto.io>
# Product: AXTO Legal — Enterprise AI Legal & Compliance Platform
# ==============================================================================
"""
Document Intelligence Engine
Upload, extract, analyze, and export legal documents.
Supports: PDF, DOCX, DOC, TXT, RTF, HTML
Output: PDF analysis reports, DOCX exports, JSON
"""
import hashlib
import io
import json
import logging
import os
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

logger = logging.getLogger("axto.legal.documents")

ALLOWED_TYPES = {
    "application/pdf":                                      "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword":                                   "doc",
    "text/plain":                                           "txt",
    "text/html":                                            "html",
    "application/rtf":                                      "rtf",
    "text/rtf":                                             "rtf",
}
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".html", ".htm", ".rtf", ".md"}
MAX_TEXT_CHARS = 50000    # ~12,500 tokens for LLM


def _doc_id(content: bytes) -> str:
    """Stable document ID from content hash."""
    return hashlib.sha256(content).hexdigest()[:24]


def _safe_filename(name: str) -> str:
    """Sanitize filename for safe storage."""
    import re
    name = re.sub(r"[^\w\s\-.]", "_", name)
    return name[:200]


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from PDF with multiple fallback strategies."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(f"[Page {i+1}]\n{text}")
            except Exception:
                pass
        return "\n\n".join(pages) if pages else ""
    except ImportError:
        pass
    except Exception as e:
        logger.warning("pypdf extraction failed: %s", e)

    # Fallback: try pdfminer
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        return pdfminer_extract(io.BytesIO(content))
    except Exception:
        pass

    return ""


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from DOCX preserving structure."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        sections = []
        for para in doc.paragraphs:
            if para.text.strip():
                # Detect headings for structure
                if para.style.name.startswith("Heading"):
                    sections.append(f"\n## {para.text}\n")
                else:
                    sections.append(para.text)
        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells)
                if row_text.strip():
                    rows.append(row_text)
            if rows:
                sections.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")
        return "\n".join(sections)
    except Exception as e:
        logger.error("DOCX extraction failed: %s", e)
        return ""


def extract_text(filename: str, content: bytes, content_type: str = "") -> tuple[str, dict]:
    """
    Extract text from document. Returns (text, metadata).
    metadata includes: page_count, word_count, language_detected
    """
    ext = Path(filename.lower()).suffix

    if ext == ".pdf" or "pdf" in content_type:
        text = extract_text_from_pdf(content)
        # Count pages
        try:
            import pypdf
            page_count = len(pypdf.PdfReader(io.BytesIO(content)).pages)
        except Exception:
            page_count = text.count("[Page ") or 1
        meta = {"format": "pdf", "page_count": page_count}

    elif ext in (".docx",) or "wordprocessingml" in content_type:
        text = extract_text_from_docx(content)
        meta = {"format": "docx", "page_count": max(1, len(text) // 3000)}

    elif ext in (".txt", ".md"):
        text = content.decode("utf-8", errors="replace")
        meta = {"format": "txt", "page_count": 1}

    elif ext in (".html", ".htm"):
        try:
            from html.parser import HTMLParser
            class _P(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.chunks = []
                def handle_data(self, d):
                    stripped = d.strip()
                    if stripped:
                        self.chunks.append(stripped)
            p = _P()
            p.feed(content.decode("utf-8", errors="replace"))
            text = " ".join(p.chunks)
        except Exception:
            text = content.decode("utf-8", errors="replace")
        meta = {"format": "html", "page_count": 1}

    else:
        text = content.decode("utf-8", errors="replace")
        meta = {"format": "unknown", "page_count": 1}

    meta["word_count"] = len(text.split())
    meta["char_count"] = len(text)
    meta["doc_id"]     = _doc_id(content)

    return text, meta


def truncate_for_llm(text: str, max_chars: int = MAX_TEXT_CHARS) -> tuple[str, bool]:
    """Truncate long documents for LLM, returning (truncated_text, was_truncated)."""
    if len(text) <= max_chars:
        return text, False
    # Smart truncation: keep beginning and end
    half = max_chars // 2
    truncated = (
        text[:half]
        + "\n\n[... DOCUMENT TRUNCATED — middle section omitted to fit LLM context ...]\n\n"
        + text[-half:]
    )
    return truncated, True


def save_upload(content: bytes, filename: str, upload_path: str) -> dict:
    """Save uploaded file and return metadata."""
    safe_name = _safe_filename(filename)
    doc_id    = _doc_id(content)
    ext       = Path(safe_name).suffix.lower()
    dest_name = f"{doc_id}_{safe_name}"
    dest_path = Path(upload_path) / dest_name

    Path(upload_path).mkdir(parents=True, exist_ok=True)
    dest_path.write_bytes(content)

    return {
        "doc_id":    doc_id,
        "filename":  safe_name,
        "dest_path": str(dest_path),
        "size_kb":   len(content) // 1024,
        "uploaded_at": datetime.now(timezone.utc).isoformat(),
    }


# ── PDF Report Generation ─────────────────────────────────────────────────────
def generate_pdf_report(
    title: str,
    analysis: dict,
    doc_meta: dict,
    output_path: str,
    language: str = "en",
) -> bytes:
    """
    Generate a professional PDF analysis report.
    Returns PDF bytes for download.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor, black, white
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            HRFlowable, KeepTogether,
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf, pagesize=A4,
            leftMargin=20*mm, rightMargin=20*mm,
            topMargin=20*mm, bottomMargin=20*mm,
        )

        PRIMARY   = HexColor("#0d1f3c")
        ACCENT    = HexColor("#0d9488")
        DANGER    = HexColor("#dc2626")
        WARNING   = HexColor("#d97706")
        SUCCESS   = HexColor("#16a34a")
        LIGHT_BG  = HexColor("#f8fafc")
        BORDER    = HexColor("#e2e8f0")

        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("h1", parent=styles["Heading1"],
                             fontSize=20, textColor=PRIMARY, spaceAfter=6,
                             fontName="Helvetica-Bold")
        h2 = ParagraphStyle("h2", parent=styles["Heading2"],
                             fontSize=13, textColor=PRIMARY, spaceAfter=4,
                             fontName="Helvetica-Bold", spaceBefore=10)
        h3 = ParagraphStyle("h3", parent=styles["Heading3"],
                             fontSize=11, textColor=ACCENT, spaceAfter=3,
                             fontName="Helvetica-Bold")
        body = ParagraphStyle("body", parent=styles["Normal"],
                              fontSize=9, textColor=HexColor("#374151"),
                              spaceAfter=4, leading=14, alignment=TA_JUSTIFY)
        small = ParagraphStyle("small", parent=styles["Normal"],
                               fontSize=7.5, textColor=HexColor("#6b7280"),
                               spaceAfter=2, leading=11)
        caption = ParagraphStyle("caption", parent=styles["Normal"],
                                  fontSize=8, textColor=HexColor("#6b7280"),
                                  alignment=TA_CENTER)

        story = []

        # ── Cover header ──────────────────────────────────────────────────
        header_data = [[
            Paragraph(f"<b>⚖️ AXTO Legal</b>", ParagraphStyle("hd", fontSize=9, textColor=white, fontName="Helvetica-Bold")),
            Paragraph(f"axto.io/legal", ParagraphStyle("hd2", fontSize=9, textColor=HexColor("#94a3b8"), alignment=2)),
        ]]
        header_tbl = Table(header_data, colWidths=["70%", "30%"])
        header_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), PRIMARY),
            ("PADDING",    (0,0), (-1,-1), 8),
            ("VALIGN",     (0,0), (-1,-1), "MIDDLE"),
        ]))
        story.append(header_tbl)
        story.append(Spacer(1, 8*mm))

        # Title
        story.append(Paragraph(title, h1))
        story.append(Paragraph(
            f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')} | "
            f"Document: {doc_meta.get('filename', '')} | "
            f"Pages: {doc_meta.get('page_count', '?')} | "
            f"Words: {doc_meta.get('word_count', '?'):,}",
            small
        ))
        story.append(HRFlowable(width="100%", thickness=2, color=ACCENT))
        story.append(Spacer(1, 5*mm))

        def _add_section(heading: str, content):
            story.append(Paragraph(heading, h2))
            if isinstance(content, str):
                story.append(Paragraph(content or "N/A", body))
            elif isinstance(content, list):
                for item in content:
                    if isinstance(item, dict):
                        for k, v in item.items():
                            story.append(Paragraph(f"<b>{k}:</b> {v}", body))
                        story.append(Spacer(1, 2*mm))
                    else:
                        story.append(Paragraph(f"• {item}", body))
            elif isinstance(content, dict):
                rows = [[str(k), str(v)] for k, v in content.items()]
                if rows:
                    tbl = Table(rows, colWidths=["35%", "65%"])
                    tbl.setStyle(TableStyle([
                        ("BACKGROUND", (0,0), (0,-1), LIGHT_BG),
                        ("FONTNAME",   (0,0), (0,-1), "Helvetica-Bold"),
                        ("FONTSIZE",   (0,0), (-1,-1), 8.5),
                        ("GRID",       (0,0), (-1,-1), 0.5, BORDER),
                        ("PADDING",    (0,0), (-1,-1), 5),
                        ("VALIGN",     (0,0), (-1,-1), "TOP"),
                    ]))
                    story.append(tbl)
            story.append(Spacer(1, 3*mm))

        # ── Analysis sections ─────────────────────────────────────────────
        if "summary" in analysis:
            _add_section("Executive Summary", analysis["summary"])
        if "contract_type" in analysis:
            _add_section("Contract Type", analysis["contract_type"])
        if "parties" in analysis:
            _add_section("Parties Involved", analysis["parties"])
        if "key_clauses" in analysis:
            _add_section("Key Clauses", analysis["key_clauses"])
        if "high_risk_clauses" in analysis:
            story.append(Paragraph("High-Risk Clauses", h2))
            risks = analysis["high_risk_clauses"]
            if isinstance(risks, list):
                for risk in risks:
                    level = str(risk.get("risk_level", "HIGH")).upper()
                    color = DANGER if "CRITICAL" in level or "HIGH" in level else (WARNING if "MEDIUM" in level else SUCCESS)
                    risk_data = [[
                        Paragraph(f"<b>[{level}]</b>", ParagraphStyle("rl", fontSize=9, textColor=color, fontName="Helvetica-Bold")),
                        Paragraph(str(risk.get("clause", risk.get("text", ""))), body),
                    ]]
                    risk_tbl = Table(risk_data, colWidths=["15%", "85%"])
                    risk_tbl.setStyle(TableStyle([
                        ("VALIGN", (0,0), (-1,-1), "TOP"),
                        ("LEFTPADDING", (0,0), (-1,-1), 4),
                    ]))
                    story.append(risk_tbl)
                    if risk.get("recommendation"):
                        story.append(Paragraph(f"→ {risk['recommendation']}", small))
                    story.append(Spacer(1, 2*mm))
        if "missing_clauses" in analysis:
            _add_section("Missing / Recommended Clauses", analysis["missing_clauses"])
        if "jurisdiction" in analysis:
            _add_section("Jurisdiction & Governing Law", analysis["jurisdiction"])
        if "overall_risk" in analysis:
            _add_section("Overall Risk Assessment", str(analysis["overall_risk"]))
        if "recommended_actions" in analysis:
            _add_section("Recommended Actions", analysis["recommended_actions"])

        # Render any other fields not handled above
        skip_keys = {"summary","contract_type","parties","key_clauses","high_risk_clauses",
                     "missing_clauses","jurisdiction","overall_risk","recommended_actions",
                     "meta","_parse_error","raw"}
        for k, v in analysis.items():
            if k not in skip_keys and v:
                _add_section(k.replace("_", " ").title(), v)

        # ── Footer ────────────────────────────────────────────────────────
        story.append(Spacer(1, 8*mm))
        story.append(HRFlowable(width="100%", thickness=1, color=BORDER))
        story.append(Spacer(1, 3*mm))
        story.append(Paragraph(
            "⚠️  DISCLAIMER: This report is generated by AXTO Legal AI and provides legal information, "
            "not legal advice. Always consult qualified legal counsel before making legal decisions. "
            "Copyright © 2024-2026 Yusron Efendi. All rights reserved. AXTO® (axto.io).",
            small
        ))

        doc.build(story)
        return buf.getvalue()

    except ImportError:
        # Fallback: plain text report as bytes
        logger.warning("reportlab not installed, generating plain text report")
        lines = [
            f"AXTO LEGAL — ANALYSIS REPORT",
            f"Copyright (c) 2024-2026 Yusron Efendi. axto.io",
            f"{'='*60}",
            f"Title: {title}",
            f"Document: {doc_meta.get('filename', '')}",
            f"Generated: {datetime.now(timezone.utc).isoformat()}",
            f"{'='*60}",
            "",
        ]
        for k, v in analysis.items():
            if k.startswith("_"):
                continue
            lines.append(f"\n{k.upper().replace('_', ' ')}")
            lines.append("-" * 40)
            lines.append(json.dumps(v, indent=2, ensure_ascii=False) if not isinstance(v, str) else v)
        lines.append("\n" + "="*60)
        lines.append("DISCLAIMER: Legal information only. Consult qualified counsel.")
        return "\n".join(lines).encode("utf-8")


def generate_docx_report(title: str, analysis: dict, doc_meta: dict) -> bytes:
    """Generate DOCX format analysis report for download."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = DocxDocument()

        # Set margins
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1.2)
            section.right_margin  = Inches(1.2)

        # Header
        hdr = doc.add_heading("AXTO LEGAL — ANALYSIS REPORT", 0)
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = meta_p.add_run(
            f"Title: {title}\n"
            f"Document: {doc_meta.get('filename', '')} | "
            f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}\n"
            f"Copyright © 2024-2026 Yusron Efendi | axto.io"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

        doc.add_paragraph()

        def add_section(heading: str, content):
            h = doc.add_heading(heading, 2)
            if isinstance(content, str):
                doc.add_paragraph(content or "N/A")
            elif isinstance(content, list):
                for item in content:
                    p = doc.add_paragraph(style="List Bullet")
                    if isinstance(item, dict):
                        p.add_run(json.dumps(item, ensure_ascii=False))
                    else:
                        p.add_run(str(item))
            elif isinstance(content, dict):
                for k, v in content.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{k}: ").bold = True
                    p.add_run(str(v))

        for k, v in analysis.items():
            if k.startswith("_") or not v:
                continue
            add_section(k.replace("_", " ").title(), v)

        doc.add_paragraph()
        disclaimer = doc.add_paragraph(
            "DISCLAIMER: This report provides legal information, not legal advice. "
            "Consult qualified legal counsel before making legal decisions. "
            "Copyright © 2024-2026 Yusron Efendi. AXTO® (axto.io)"
        )
        disclaimer.runs[0].font.size = Pt(8)
        disclaimer.runs[0].font.color.rgb = RGBColor(0x94, 0xa3, 0xb8)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        # Fallback: just return the JSON analysis as text
        return json.dumps(analysis, indent=2, ensure_ascii=False).encode("utf-8")
